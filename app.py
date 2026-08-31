import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import matplotlib.transforms as transforms
import numpy as np
import os
import io
import docx
from docx.shared import Inches
from datetime import datetime
import scipy.stats as stats  

# =========================
# PAGE CONFIG
# =========================
st.set_page_config(
    page_title="Gloss Analysis Dashboard",
    page_icon="📊",
    layout="wide"
)

# View 1 Background CSS
st.markdown(
    """
    <style>
    .stApp {
        background: linear-gradient(270deg, #ffffff, #f0f9ff, #e0f2fe, #fef3c7, #ecfeff);
        background-size: 800% 800%;
        animation: gradientBG 20s ease infinite;
    }
    @keyframes gradientBG {
        0% { background-position: 0% 50%; }
        50% { background-position: 100% 50%; }
        100% { background-position: 0% 50%; }
    }
    [data-testid="stSidebar"] {background-color: #f6f8fa;}
    </style>
    """,
    unsafe_allow_html=True
)

# =========================
# LOAD DATA & CLEANING
# =========================
DIR_PATH = r"D:\Mandy\Source data\统计室内用途管控膜厚色差\光泽\2026\gloss 2526"
PROD_FILE = os.path.join(DIR_PATH, "Production_Data_Merge_Result.xlsx")
LAB_FILE = os.path.join(DIR_PATH, "塗料檢驗報告.xlsx")

@st.cache_data(ttl=300)
def load_datasets():
    prod_df = pd.DataFrame()
    if os.path.exists(PROD_FILE):
        prod_df = pd.read_excel(PROD_FILE)
    elif os.path.exists(PROD_FILE.replace('.xlsx', '.csv')):
        prod_df = pd.read_csv(PROD_FILE.replace('.xlsx', '.csv'))
        
    lab_df = pd.DataFrame()
    if os.path.exists(LAB_FILE):
        lab_df = pd.read_excel(LAB_FILE)
    elif os.path.exists(LAB_FILE.replace('.xlsx', '.csv')):
        lab_df = pd.read_csv(LAB_FILE.replace('.xlsx', '.csv'))
        
    return prod_df, lab_df

df_raw, df_lab_raw = load_datasets()

st.title("📊 Gloss SPC Dashboard (LAB vs LINE)")

if not df_raw.empty:
    # ---------------------------------------------------------
    # DATA PREPARATION & MAPPING
    # ---------------------------------------------------------
    coil_col = "產出鋼捲號碼"
    lab_gloss_col = "光澤"
    line_north_col = "NORTH_TOP_BLANCH"
    line_south_col = "SOUTH_TOP_BLANCH"
    paint_code_col = "面漆代號"
    date_col = "生產日期"
    batch_col = "TOPPAINT_BATCH_NO"
    
    required_cols = [coil_col, lab_gloss_col, line_north_col, line_south_col, date_col, paint_code_col]
    missing_cols = [col for col in required_cols if col not in df_raw.columns]
    
    if missing_cols:
        st.error(f"Missing required columns in dataset: {missing_cols}")
    else:
        if batch_col not in df_raw.columns:
            batch_col = "訂單號碼"

        # 1. Drop true NaNs in critical columns first
        df = df_raw.dropna(subset=[coil_col, date_col, batch_col]).copy()
        
        # 2. Convert to string and strip whitespaces
        df[coil_col] = df[coil_col].astype(str).str.strip()
        df[batch_col] = df[batch_col].astype(str).str.strip()
        
        # 3. Strict filtering: Remove literal "nan", "none", "na", or empty strings
        invalid_strs = ["nan", "none", "na", "null", ""]
        df = df[~df[coil_col].str.lower().isin(invalid_strs)]
        df = df[~df[batch_col].str.lower().isin(invalid_strs)]
        
        # 4. Convert gloss values to numeric
        df[lab_gloss_col] = pd.to_numeric(df[lab_gloss_col], errors='coerce')
        df[line_north_col] = pd.to_numeric(df[line_north_col], errors='coerce')
        df[line_south_col] = pd.to_numeric(df[line_south_col], errors='coerce')
        
        df['LINE_Gloss'] = df[[line_north_col, line_south_col]].mean(axis=1)
        
        # 5. Final drop of invalid rows
        df = df.dropna(subset=[lab_gloss_col, 'LINE_Gloss'], how='all')
        df[date_col] = pd.to_datetime(df[date_col], errors='coerce')
        df = df.dropna(subset=[date_col]).reset_index(drop=True)

        # =========================================================
        # UI: SIDEBAR CONTROLS
        # =========================================================
        st.sidebar.title("🎨 Filters & Settings")
        
        app_mode = st.sidebar.radio("👁️ Select View Mode:", ["📈 Control Charts (OOC)", "📊 Distribution & Bias Analysis", "🎯 Task 3 - LAB Input Limit Optimization"])
        st.sidebar.markdown("---")

        paint_codes = sorted(df[paint_code_col].dropna().unique())
        selected_paint = st.sidebar.selectbox("Select Topcoat Code", paint_codes)
        
        df_filtered = df[df[paint_code_col] == selected_paint].copy()

        if df_filtered.empty:
            st.warning(f"No data available for topcoat code: {selected_paint}")
        else:
            if not df_lab_raw.empty and "製造批號" in df_lab_raw.columns and "檢驗日期" in df_lab_raw.columns:
                lab_subset = df_lab_raw[["製造批號", "檢驗日期"]].dropna().copy()
                lab_subset["製造批號"] = lab_subset["製造批號"].astype(str).str.strip()
                lab_subset["檢驗日期"] = pd.to_datetime(lab_subset["檢驗日期"], errors='coerce')
                lab_subset = lab_subset.sort_values("檢驗日期").drop_duplicates(subset=["製造批號"], keep="first")
                
                batch_date_map = dict(zip(lab_subset["製造批號"], lab_subset["檢驗日期"]))
                df_filtered['Batch_Input_Date'] = df_filtered[batch_col].map(batch_date_map)
            else:
                df_filtered['Batch_Input_Date'] = pd.NaT

            fallback_dates = df_filtered.groupby(batch_col)[date_col].min().to_dict()
            df_filtered['Batch_Input_Date'] = df_filtered['Batch_Input_Date'].fillna(df_filtered[batch_col].map(fallback_dates))
            df_filtered['Batch_Input_Date'] = pd.to_datetime(df_filtered['Batch_Input_Date'])

            # ---------------------------------------------------------
            # OPTIONAL PHASE II ENABLE/DISABLE
            # ---------------------------------------------------------
            st.sidebar.markdown("### ⏱️ Control Implementation")
            enable_phase2 = st.sidebar.checkbox("Enable Phase II (New Control Limits)", value=False)
            
            if enable_phase2:
                default_date = df_filtered['Batch_Input_Date'].iloc[len(df_filtered)//2].date()
                selected_date = st.sidebar.date_input("Cutoff Date (Phase II)", value=default_date)
                cutoff_date = pd.to_datetime(selected_date)
                
                df_filtered['Is_Phase_II'] = df_filtered['Batch_Input_Date'] >= cutoff_date
                df_filtered = df_filtered.sort_values(by=['Is_Phase_II', date_col, coil_col]).reset_index(drop=True)
                
                phase2_data = df_filtered[df_filtered['Is_Phase_II'] == True]
                if not phase2_data.empty:
                    control_index = phase2_data.index[0]
                else:
                    control_index = len(df_filtered)
            else:
                df_filtered['Is_Phase_II'] = False
                df_filtered = df_filtered.sort_values(by=[date_col, coil_col]).reset_index(drop=True)
                control_index = len(df_filtered)

            default_lab_mean = df_filtered[lab_gloss_col].mean()
            default_line_mean = df_filtered['LINE_Gloss'].mean()
            
            # =========================================================
            # DÒNG THỐNG KÊ THỜI GIAN VÀ SẢN LƯỢNG
            # =========================================================
            min_date = df_filtered[date_col].min().strftime('%Y-%m-%d')
            max_date = df_filtered[date_col].max().strftime('%Y-%m-%d')
            num_batches = df_filtered[batch_col].nunique()
            num_coils = df_filtered[coil_col].dropna().nunique()
            
            st.success(f"📅 **Timeframe:** {min_date} to {max_date} | **Volume:** {num_batches} Batches ({num_coils} Coils).")
            # =========================================================

            st.sidebar.markdown("### 🎛️ Control Limits Setup")
            
            with st.sidebar.expander("🧪 LAB Limits", expanded=True):
                l_b1, l_b2 = st.columns(2)
                lab_before_lcl = l_b1.number_input("LAB LCL (Phase I)", value=float(default_lab_mean - 5), step=1.0)
                lab_before_ucl = l_b2.number_input("LAB UCL (Phase I)", value=float(default_lab_mean + 5), step=1.0)
                
                if enable_phase2:
                    l_a1, l_a2 = st.columns(2)
                    lab_after_lcl = l_a1.number_input("LAB LCL (Phase II)", value=float(default_lab_mean - 4), step=1.0)
                    lab_after_ucl = l_a2.number_input("LAB UCL (Phase II)", value=float(default_lab_mean + 4), step=1.0)
                else:
                    lab_after_lcl = lab_before_lcl
                    lab_after_ucl = lab_before_ucl

            with st.sidebar.expander("🏭 LINE Limits", expanded=True):
                ln_b1, ln_b2 = st.columns(2)
                line_before_lcl = ln_b1.number_input("LINE LCL (Phase I)", value=float(default_line_mean - 5), step=1.0)
                line_before_ucl = ln_b2.number_input("LINE UCL (Phase I)", value=float(default_line_mean + 5), step=1.0)
                
                if enable_phase2:
                    ln_a1, ln_a2 = st.columns(2)
                    line_after_lcl = ln_a1.number_input("LINE LCL (Phase II)", value=float(default_line_mean - 4), step=1.0)
                    line_after_ucl = ln_a2.number_input("LINE UCL (Phase II)", value=float(default_line_mean + 4), step=1.0)
                else:
                    line_after_lcl = line_before_lcl
                    line_after_ucl = line_before_ucl

            # =========================================================
            # VIEW 1: CONTROL CHARTS (OOC)
            # =========================================================
            if app_mode == "📈 Control Charts (OOC)":
                def plot_control_chart(x_labels, lab_y, line_y, ctrl_idx, title, x_label_name, use_seq_labels=True):
                    fig, ax = plt.subplots(figsize=(12, 5))
                    ax.set_facecolor('#f2f2f2')
                    total_points = len(x_labels)
                    
                    if enable_phase2:
                        padding = max(7, int(total_points * 0.075))
                    else:
                        padding = 0
                        
                    xlim_max = total_points + padding - 0.5
                    extended_x_vals = np.arange(total_points + padding)
                    
                    ax.plot(np.arange(total_points), lab_y, marker="^", color="#548235", linestyle="-", linewidth=1.5, markersize=8, label="LAB Input")
                    ax.plot(np.arange(total_points), line_y, marker="o", color="#ffc000", linestyle="-", linewidth=1.5, markersize=8, markerfacecolor="white", markeredgewidth=2, label="LINE Output")

                    lab_before = lab_y[:ctrl_idx]
                    if len(lab_before) > 0:
                        out_lab_before = (lab_before > lab_before_ucl) | (lab_before < lab_before_lcl)
                        if out_lab_before.any():
                            ax.plot(np.arange(ctrl_idx)[out_lab_before], lab_before[out_lab_before], marker="^", color="red", linestyle="None", markersize=10, zorder=5)
                    
                    if enable_phase2:
                        lab_after = lab_y[ctrl_idx:]
                        if len(lab_after) > 0:
                            out_lab_after = (lab_after > lab_after_ucl) | (lab_after < lab_after_lcl)
                            if out_lab_after.any():
                                ax.plot(np.arange(total_points)[ctrl_idx:][out_lab_after], lab_after[out_lab_after], marker="^", color="red", linestyle="None", markersize=10, zorder=5)

                    line_before = line_y[:ctrl_idx]
                    if len(line_before) > 0:
                        out_line_before = (line_before > line_before_ucl) | (line_before < line_before_lcl)
                        if out_line_before.any():
                            ax.plot(np.arange(ctrl_idx)[out_line_before], line_before[out_line_before], marker="o", color="red", linestyle="None", markersize=10, zorder=5)
                    
                    if enable_phase2:
                        line_after = line_y[ctrl_idx:]
                        if len(line_after) > 0:
                            out_line_after = (line_after > line_after_ucl) | (line_after < line_after_lcl)
                            if out_line_after.any():
                                ax.plot(np.arange(total_points)[ctrl_idx:][out_line_after], line_after[out_line_after], marker="o", color="red", linestyle="None", markersize=10, zorder=5)

                    phase1_end = (ctrl_idx - 0.5) if (enable_phase2 and ctrl_idx < total_points) else xlim_max

                    if ctrl_idx > 0:
                        ax.hlines(lab_before_lcl, xmin=-0.5, xmax=phase1_end, colors='#7030a0', linestyles='solid', lw=2)
                        ax.hlines(lab_before_ucl, xmin=-0.5, xmax=phase1_end, colors='#7030a0', linestyles='solid', lw=2, label="LAB Limits")
                        ax.hlines(line_before_lcl, xmin=-0.5, xmax=phase1_end, colors='deepskyblue', linestyles='dashed', lw=2)
                        ax.hlines(line_before_ucl, xmin=-0.5, xmax=phase1_end, colors='deepskyblue', linestyles='dashed', lw=2, label="LINE Limits")
                    
                    if enable_phase2 and ctrl_idx < total_points:
                        start_x = ctrl_idx - 0.5 if ctrl_idx > 0 else -0.5
                        ax.hlines(lab_after_lcl, xmin=start_x, xmax=xlim_max, colors='#7030a0', linestyles='solid', lw=2)
                        ax.hlines(lab_after_ucl, xmin=start_x, xmax=xlim_max, colors='#7030a0', linestyles='solid', lw=2)
                        ax.hlines(line_after_lcl, xmin=start_x, xmax=xlim_max, colors='deepskyblue', linestyles='dashed', lw=2)
                        ax.hlines(line_after_ucl, xmin=start_x, xmax=xlim_max, colors='deepskyblue', linestyles='dashed', lw=2)
                        
                        if 0 < ctrl_idx < total_points:
                            ax.vlines(ctrl_idx - 0.5, ymin=lab_before_lcl, ymax=lab_after_lcl, colors='#7030a0', linestyles='dotted', lw=1.5)
                            ax.vlines(ctrl_idx - 0.5, ymin=lab_before_ucl, ymax=lab_after_ucl, colors='#7030a0', linestyles='dotted', lw=1.5)
                            ax.vlines(ctrl_idx - 0.5, ymin=line_before_lcl, ymax=line_after_lcl, colors='deepskyblue', linestyles='dotted', lw=1.5)
                            ax.vlines(ctrl_idx - 0.5, ymin=line_before_ucl, ymax=line_after_ucl, colors='deepskyblue', linestyles='dotted', lw=1.5)

                    trans = transforms.blended_transform_factory(ax.transData, ax.transAxes)
                    if enable_phase2 and 0 < ctrl_idx < total_points:
                        ax.axvline(x=ctrl_idx - 0.5, color="#000000", linestyle=(0, (3, 3)), linewidth=1.5)
                        ax.text(ctrl_idx - 0.5, 1.02, "  After Control", color="#0070c0", fontsize=12, ha="left", va="bottom", transform=trans)
                        ax.text(ctrl_idx - 0.5, 1.02, "Before Control  ", color="#0070c0", fontsize=12, ha="right", va="bottom", transform=trans)

                    ax.set_xlim(-0.5, xlim_max)
                    ax.set_xticks(extended_x_vals)
                    
                    if use_seq_labels:
                        final_labels = [str(i + 1) if i < total_points else "" for i in range(total_points + padding)]
                        ax.set_xticklabels(final_labels, rotation=0, ha='center')
                    else:
                        final_labels = [str(x_labels[i]) if i < total_points else "" for i in range(total_points + padding)]
                        ax.set_xticklabels(final_labels, rotation=45, ha='right')
                    
                    if len(extended_x_vals) > 30 and use_seq_labels:
                        step = max(1, len(extended_x_vals) // 20)
                        for i, label in enumerate(ax.xaxis.get_ticklabels()):
                            if i % step != 0: label.set_visible(False)

                    ax.set_title(title, fontsize=15, fontweight="bold", pad=25)
                    ax.set_xlabel(x_label_name)
                    ax.set_ylabel("Gloss Value")
                    ax.grid(axis="y", color="#cccccc", linestyle="-", linewidth=1)
                    ax.grid(axis="x", visible=False)
                    ax.legend(bbox_to_anchor=(1.02, 1), loc="upper left", frameon=True, edgecolor="black")
                    fig.subplots_adjust(right=0.8, top=0.8, bottom=0.2)
                    return fig

                st.markdown(f"### 📈 LAB vs LINE Gloss Control Chart: {selected_paint} (Coil Level)")

                # -----------------------------------------------------
                # COIL-LEVEL DATA: 1 output coil = 1 observation
                # This prevents duplicate records of the same 產出鋼捲號碼
                # from being plotted multiple times.
                # -----------------------------------------------------
                df_coil = (
                    df_filtered
                    .groupby(coil_col, as_index=False)
                    .agg({
                        date_col: 'min',
                        'Batch_Input_Date': 'first',
                        batch_col: 'first',
                        'Is_Phase_II': 'first',
                        lab_gloss_col: 'mean',
                        line_north_col: 'mean',
                        line_south_col: 'mean',
                        'LINE_Gloss': 'mean'
                    })
                )

                if enable_phase2:
                    df_coil = (
                        df_coil
                        .sort_values(by=['Is_Phase_II', date_col, coil_col])
                        .reset_index(drop=True)
                    )
                    phase2_coil_data = df_coil[df_coil['Is_Phase_II'] == True]
                    control_index_coil = (
                        phase2_coil_data.index[0]
                        if not phase2_coil_data.empty
                        else len(df_coil)
                    )
                else:
                    df_coil = (
                        df_coil
                        .sort_values(by=[date_col, coil_col])
                        .reset_index(drop=True)
                    )
                    control_index_coil = len(df_coil)

                fig_coil = plot_control_chart(
                    df_coil[coil_col].tolist(),
                    df_coil[lab_gloss_col].values,
                    df_coil['LINE_Gloss'].values,
                    control_index_coil,
                    f"COMBINED GLOSS PROCESS: {selected_paint} (By Coil)",
                    "Sequential Coil Count",
                    use_seq_labels=True
                )
                st.pyplot(fig_coil)

                st.caption(
                    f"Coil-level chart uses **1 產出鋼捲號碼 = 1 observation** "
                    f"({len(df_coil)} unique coils)."
                )

                st.markdown("---")
                st.markdown(f"### 📊 LAB vs LINE Gloss Control Chart: {selected_paint} (Aggregated by Batch)")
                
                df_batch = df_filtered.groupby(batch_col, as_index=False).agg({
                    'Batch_Input_Date': 'first', 'Is_Phase_II': 'first', lab_gloss_col: 'mean', 'LINE_Gloss': 'mean'
                }).sort_values(by=['Is_Phase_II', 'Batch_Input_Date']).reset_index(drop=True)
                
                if enable_phase2:
                    phase2_batch_data = df_batch[df_batch['Is_Phase_II'] == True]
                    control_index_batch = phase2_batch_data.index[0] if not phase2_batch_data.empty else len(df_batch)
                else:
                    control_index_batch = len(df_batch)

                fig_batch = plot_control_chart(
                    df_batch[batch_col].tolist(), df_batch[lab_gloss_col].values, df_batch['LINE_Gloss'].values,
                    control_index_batch, f"COMBINED GLOSS PROCESS: {selected_paint} (By Batch)", "Batch Number", use_seq_labels=False
                )
                st.pyplot(fig_batch)

                # Word Report Generator for Control Charts
                st.markdown("---")
                st.markdown("### 📄 Export Technical Report")
                def generate_word_report(fig1, fig2, paint_code):
                    doc = docx.Document()
                    doc.add_heading(f'光澤 SPC 分析報告 (Gloss SPC Analysis Report) - {paint_code}', 0)
                    doc.add_paragraph(f'產生時間 (Generated on): {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
                    
                    doc.add_heading('1. 鋼捲級別管制圖 (Coil Level Control Chart)', level=1)
                    buf_coil = io.BytesIO()
                    fig1.savefig(buf_coil, format='png', dpi=150, bbox_inches='tight')
                    buf_coil.seek(0)
                    doc.add_picture(buf_coil, width=Inches(6.0))
                    
                    doc.add_heading('2. 批次級別管制圖 (Batch Level Control Chart)', level=1)
                    buf_batch = io.BytesIO()
                    fig2.savefig(buf_batch, format='png', dpi=150, bbox_inches='tight')
                    buf_batch.seek(0)
                    doc.add_picture(buf_batch, width=Inches(6.0))
                    
                    doc_buf = io.BytesIO()
                    doc.save(doc_buf)
                    doc_buf.seek(0)
                    return doc_buf

                word_buffer = generate_word_report(fig_coil, fig_batch, selected_paint)
                st.download_button(
                    label="📥 Download Word Report (.docx)",
                    data=word_buffer,
                    file_name=f"Gloss_SPC_Report_{selected_paint}_{datetime.now().strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

            # =========================================================
            # VIEW 2: DISTRIBUTION & BIAS ANALYSIS
            # =========================================================
            elif app_mode == "📊 Distribution & Bias Analysis":
                st.markdown(f"### 📊 Gloss Distribution & Bias Analysis: {selected_paint}")
                
                lab_data = df_filtered[lab_gloss_col].dropna()
                line_data = df_filtered['LINE_Gloss'].dropna()
                
                if lab_data.empty or line_data.empty:
                    st.warning("Not enough data to generate distribution.")
                else:
                    lab_mu, lab_sigma = lab_data.mean(), lab_data.std()
                    line_mu, line_sigma = line_data.mean(), line_data.std()
                    bias = line_mu - lab_mu
                    
                    st.info(f"💡 **Bias Analysis:** LINE output is on average **{abs(bias):.2f} GU** {'higher' if bias > 0 else 'lower'} than LAB input.")

                    # =========================================================
                    # TASK 2 ADDITION:
                    # (1) GlossLine = GlossLab + ΔBias + ε
                    # (2) TargetLab = TargetLine - ΔBias
                    # (3) TargetLine = (USL + LSL) / 2
                    # (4) ICL = TargetLab ± n * σLab
                    # =========================================================
                    st.markdown("### 🎯 Bias Compensation & Input Control Limit (ICL)")

                    icl_n = st.number_input(
                        "Control Multiplier n",
                        min_value=0.1,
                        max_value=5.0,
                        value=1.0,
                        step=0.1,
                        key=f"task2_icl_n_{selected_paint}",
                        help="ε = n × σLab. Default n = 1 according to the proposed method."
                    )

                    target_line = (line_before_ucl + line_before_lcl) / 2.0
                    delta_bias = bias
                    target_lab = target_line - delta_bias
                    epsilon_icl = icl_n * lab_sigma
                    icl_lcl = target_lab - epsilon_icl
                    icl_ucl = target_lab + epsilon_icl

                    t2c1, t2c2, t2c3, t2c4, t2c5 = st.columns(5)
                    t2c1.metric("ΔBias", f"{delta_bias:+.2f} GU")
                    t2c2.metric("Target LINE", f"{target_line:.2f} GU")
                    t2c3.metric("Target LAB", f"{target_lab:.2f} GU")
                    t2c4.metric("σ LAB", f"{lab_sigma:.2f} GU")
                    t2c5.metric("Calculated ICL", f"{icl_lcl:.2f} ~ {icl_ucl:.2f} GU")

                    st.markdown(
                        f"""
                        **Calculation based on the proposed theory**

                        - (1) `GlossLine = GlossLab + ΔBias + ε`
                        - `ΔBias = Mean(LINE) − Mean(LAB) = {line_mu:.2f} − {lab_mu:.2f} = {delta_bias:+.2f} GU`
                        - (3) `TargetLine = (USL + LSL) / 2 = ({line_before_ucl:.2f} + {line_before_lcl:.2f}) / 2 = {target_line:.2f} GU`
                        - (2) `TargetLab = TargetLine − ΔBias = {target_line:.2f} − ({delta_bias:+.2f}) = {target_lab:.2f} GU`
                        - `ε = n × σLab = {icl_n:.1f} × {lab_sigma:.2f} = {epsilon_icl:.2f} GU`
                        - (4) `ICL = TargetLab ± ε = {target_lab:.2f} ± {epsilon_icl:.2f}`
                        - **Calculated ICL = {icl_lcl:.2f} ~ {icl_ucl:.2f} GU**
                        """
                    )

                    hist_lab_min = lab_data.min()
                    hist_lab_max = lab_data.max()

                    if target_lab < hist_lab_min or target_lab > hist_lab_max:
                        st.warning(
                            f"⚠️ Calculated Target LAB ({target_lab:.2f} GU) is outside the historical LAB range "
                            f"({hist_lab_min:.2f} ~ {hist_lab_max:.2f} GU). Treat it as a pilot target and validate "
                            f"before formal adoption."
                        )
                    else:
                        st.success(
                            f"✅ Calculated Target LAB ({target_lab:.2f} GU) is within the historical LAB range "
                            f"({hist_lab_min:.2f} ~ {hist_lab_max:.2f} GU)."
                        )

                    
                    fig_dist, ax_d = plt.subplots(figsize=(12, 6))
                    ax_d.set_facecolor('#ffffff')
                    
                    min_val = min(lab_data.min(), line_data.min())
                    max_val = max(lab_data.max(), line_data.max())
                    bins = np.linspace(min_val - 1, max_val + 1, 15)
                    bin_width = bins[1] - bins[0]
                    
                    # 1. Plot Histograms
                    ax_d.hist(lab_data, bins=bins, alpha=0.4, color='tab:blue', label='Lab Histogram', edgecolor='white', linewidth=1.2)
                    ax_d.hist(line_data, bins=bins, alpha=0.4, color='tab:orange', label='Line Histogram', edgecolor='white', linewidth=1.2)
                    
                    # 2. Plot Normal Distribution Curves 
                    x_axis = np.linspace(min_val - 3, max_val + 3, 500)
                    
                    if lab_sigma > 0:
                        lab_pdf = stats.norm.pdf(x_axis, lab_mu, lab_sigma) * len(lab_data) * bin_width
                        ax_d.plot(x_axis, lab_pdf, color='tab:blue', lw=2.5, label=f'Lab Curve (σ={lab_sigma:.2f})')
                    
                    if line_sigma > 0:
                        line_pdf = stats.norm.pdf(x_axis, line_mu, line_sigma) * len(line_data) * bin_width
                        ax_d.plot(x_axis, line_pdf, color='tab:orange', lw=2.5, label=f'Line Curve (σ={line_sigma:.2f})')
                        
                    y_max_current = ax_d.get_ylim()[1]
                    ax_d.set_ylim(0, y_max_current * 1.3)
                    y_max = y_max_current 
                    
                    # 3. Draw Mean Lines & Non-overlapping Labels
                    # Fixed vertical rows in axes coordinates prevent labels from covering each other.
                    ax_d.axvline(lab_mu, color='tab:blue', linestyle='--', lw=1.5)
                    props_lab = dict(boxstyle='round,pad=0.35', facecolor='tab:blue', alpha=0.92, edgecolor='none')
                    ax_d.annotate(
                        f"Lab\nμ: {lab_mu:.1f} | σ: {lab_sigma:.2f}",
                        xy=(lab_mu, 0.78),
                        xycoords=('data', 'axes fraction'),
                        color='white',
                        ha='center',
                        va='top',
                        bbox=props_lab,
                        fontweight='bold',
                        fontsize=8.5,
                        zorder=20
                    )

                    ax_d.axvline(line_mu, color='tab:orange', linestyle='--', lw=1.5)
                    props_line = dict(boxstyle='round,pad=0.35', facecolor='tab:orange', alpha=0.92, edgecolor='none')
                    ax_d.annotate(
                        f"Line\nμ: {line_mu:.1f} | σ: {line_sigma:.2f}",
                        xy=(line_mu, 0.78),
                        xycoords=('data', 'axes fraction'),
                        color='white',
                        ha='center',
                        va='top',
                        bbox=props_line,
                        fontweight='bold',
                        fontsize=8.5,
                        zorder=20
                    )

                    # 4. Draw Specification Limits & Non-overlapping Labels
                    props_lab_lsl = dict(boxstyle='round,pad=0.28', facecolor='red', alpha=0.92, edgecolor='none')
                    props_line_lsl = dict(boxstyle='round,pad=0.28', facecolor='forestgreen', alpha=0.92, edgecolor='none')

                    ax_d.axvline(lab_before_lcl, color='red', linestyle='--', lw=1.5)
                    ax_d.axvline(lab_before_ucl, color='red', linestyle='--', lw=1.5)
                    ax_d.annotate(
                        f"Lab LSL\n{lab_before_lcl:.1f}",
                        xy=(lab_before_lcl, 0.89),
                        xycoords=('data', 'axes fraction'),
                        color='white',
                        ha='center',
                        va='top',
                        bbox=props_lab_lsl,
                        fontweight='bold',
                        fontsize=8,
                        zorder=21
                    )
                    ax_d.annotate(
                        f"Lab USL\n{lab_before_ucl:.1f}",
                        xy=(lab_before_ucl, 0.89),
                        xycoords=('data', 'axes fraction'),
                        color='white',
                        ha='center',
                        va='top',
                        bbox=props_lab_lsl,
                        fontweight='bold',
                        fontsize=8,
                        zorder=21
                    )

                    ax_d.axvline(line_before_lcl, color='forestgreen', linestyle='--', lw=1.5)
                    ax_d.axvline(line_before_ucl, color='forestgreen', linestyle='--', lw=1.5)
                    ax_d.annotate(
                        f"Line LSL\n{line_before_lcl:.1f}",
                        xy=(line_before_lcl, 0.99),
                        xycoords=('data', 'axes fraction'),
                        color='white',
                        ha='center',
                        va='top',
                        bbox=props_line_lsl,
                        fontweight='bold',
                        fontsize=8,
                        zorder=22
                    )
                    ax_d.annotate(
                        f"Line USL\n{line_before_ucl:.1f}",
                        xy=(line_before_ucl, 0.99),
                        xycoords=('data', 'axes fraction'),
                        color='white',
                        ha='center',
                        va='top',
                        bbox=props_line_lsl,
                        fontweight='bold',
                        fontsize=8,
                        zorder=22
                    )

                    # Bias-compensated Target LAB and calculated ICL
                    ax_d.axvline(
                        target_lab,
                        color='black',
                        linestyle='-.',
                        lw=2.0,
                        label=f'Target LAB ({target_lab:.2f})'
                    )
                    ax_d.axvspan(
                        icl_lcl,
                        icl_ucl,
                        color='gray',
                        alpha=0.10,
                        label=f'Calculated ICL ({icl_lcl:.2f}~{icl_ucl:.2f})'
                    )

                    # 5. Styling and Labels
                    ax_d.set_xlabel("Gloss Value (GU)", fontweight="bold")
                    ax_d.set_ylabel("Number of Coils", fontweight="bold")
                    ax_d.grid(True, linestyle='-', alpha=0.5, color='#e0e0e0')
                    
                    ax_d.legend(bbox_to_anchor=(1.02, 1), loc="upper left", frameon=True, edgecolor="black", fontsize=9)
                    
                    fig_dist.subplots_adjust(right=0.80, top=0.94, bottom=0.12) 
                    st.pyplot(fig_dist)

                # Word Report Generator cho Distribution View
                st.markdown("---")
                st.markdown("### 📄 Export Technical Report")
                
                def generate_dist_word_report(fig, paint_code):
                    doc = docx.Document()
                    doc.add_heading(f'光澤 分佈與偏差分析 (Gloss Distribution & Bias Analysis) - {paint_code}', 0)
                    doc.add_paragraph(f'產生時間 (Generated on): {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
                    
                    # Thêm phân tích Bias
                    doc.add_heading('1. 統計摘要 (Statistical Summary)', level=1)
                    bias_text = f"LINE 產線的平均輸出比 LAB 檢驗輸入 {'高' if bias > 0 else '低'} {abs(bias):.2f} GU。"
                    doc.add_paragraph(bias_text)
                    
                    doc.add_heading('2. 偏差補償與入料管制界限 (Bias Compensation & ICL)', level=1)
                    doc.add_paragraph(
                        f"GlossLine = GlossLab + ΔBias + ε；ΔBias = {delta_bias:+.2f} GU。"
                    )
                    doc.add_paragraph(
                        f"TargetLine = ({line_before_ucl:.2f} + {line_before_lcl:.2f}) / 2 = {target_line:.2f} GU；"
                        f"TargetLab = {target_line:.2f} - ({delta_bias:+.2f}) = {target_lab:.2f} GU。"
                    )
                    doc.add_paragraph(
                        f"σLab = {lab_sigma:.2f} GU，n = {icl_n:.1f}，"
                        f"ε = n × σLab = {epsilon_icl:.2f} GU；"
                        f"ICL = {icl_lcl:.2f} ~ {icl_ucl:.2f} GU。"
                    )

                    doc.add_heading('3. 分佈與偏差圖 (Distribution & Bias Chart)', level=1)

                    buf = io.BytesIO()
                    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
                    buf.seek(0)
                    doc.add_picture(buf, width=Inches(6.0))
                    
                    doc_buf = io.BytesIO()
                    doc.save(doc_buf)
                    doc_buf.seek(0)
                    return doc_buf

                if not lab_data.empty and not line_data.empty:
                    word_buffer_dist = generate_dist_word_report(fig_dist, selected_paint)
                    st.download_button(
                        label="📥 Download Distribution Report (.docx)",
                        data=word_buffer_dist,
                        file_name=f"Gloss_Distribution_Report_{selected_paint}_{datetime.now().strftime('%Y%m%d')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )


            # =========================================================
            # VIEW 3 / TASK 3: LAB INPUT LIMIT OPTIMIZATION
            # THREE-LOGIC METHOD
            #   Logic 1: Historical Safety Control
            #   Logic 2: Center-to-Center Optimization
            #   Logic 3: Controlled Pilot Toward Optimized LAB Center
            # =========================================================
            elif app_mode == "🎯 Task 3 - LAB Input Limit Optimization":
                st.markdown(f"### 🎯 Task 3 - LAB Input Limit Optimization: {selected_paint}")
                st.caption(
                    "Three-logics method: establish a historically safe LAB range, calculate the exact center gap "
                    "between the current LINE output and the midpoint of the active LINE limits, then translate that "
                    "LINE gap into the LAB adjustment required to move output toward the LINE target center."
                )

                # -----------------------------------------------------
                # TASK 3 SETTINGS
                # -----------------------------------------------------
                st.sidebar.markdown("---")
                st.sidebar.markdown("### 🎯 Task 3 Settings")

                task3_min_coils = st.sidebar.number_input(
                    "Minimum Coils per Candidate Range",
                    min_value=5, max_value=500, value=20, step=1
                )

                task3_target_pass = st.sidebar.number_input(
                    "Required LINE Pass Rate (%)",
                    min_value=50.0, max_value=100.0, value=95.0, step=0.5
                )

                task3_step = st.sidebar.number_input(
                    "LAB Search Step (GU)",
                    min_value=0.1, max_value=5.0, value=0.5, step=0.1
                )

                task3_min_coverage = st.sidebar.number_input(
                    "Minimum Historical Coverage (%)",
                    min_value=10.0, max_value=100.0, value=70.0, step=5.0,
                    help="Candidate LAB range must retain at least this percentage of all historical coils."
                )

                task3_max_center_dev = st.sidebar.number_input(
                    "Maximum LINE Center Deviation (GU)",
                    min_value=0.0, max_value=20.0, value=2.0, step=0.5,
                    help=(
                        "LINE is treated as sufficiently centered when the observed LINE center is within this "
                        "distance from the midpoint of the active LINE limits."
                    )
                )

                task3_require_confidence = st.sidebar.checkbox(
                    "Require confidence check",
                    value=True,
                    help=(
                        "Candidate 95% Wilson lower pass bound must be at least as high as the current-control "
                        "lower confidence bound."
                    )
                )

                st.sidebar.markdown("#### Center Optimization / Pilot")

                task3_reverse_min_r2 = st.sidebar.number_input(
                    "Minimum R² for Model-Based Shift",
                    min_value=0.0, max_value=1.0, value=0.20, step=0.05,
                    help=(
                        "A full Required LAB Shift is calculated only when the historical LAB→LINE relationship "
                        "is sufficiently reliable."
                    )
                )

                task3_reverse_min_abs_slope = st.sidebar.number_input(
                    "Minimum |LINE/LAB Slope|",
                    min_value=0.01, max_value=10.0, value=0.10, step=0.05,
                    help="Prevents unstable reverse calculation when the fitted slope is too close to zero."
                )

                task3_max_step = st.sidebar.number_input(
                    "Maximum LAB Adjustment per Pilot (GU)",
                    min_value=0.1, max_value=10.0, value=1.0, step=0.1,
                    help="Maximum movement from the current LAB control center in one pilot cycle."
                )

                task3_pilot_half_width = st.sidebar.number_input(
                    "Pilot Range Half Width (GU)",
                    min_value=0.1, max_value=5.0, value=0.5, step=0.1
                )

                task3_max_pilot_outside_history = st.sidebar.number_input(
                    "Maximum Pilot Distance Outside History (GU)",
                    min_value=0.0, max_value=10.0, value=1.0, step=0.5,
                    help=(
                        "A pilot target cannot move farther than this distance outside the historical LAB input range."
                    )
                )

                # -----------------------------------------------------
                # ACTIVE LIMITS
                # When Phase II is enabled, Task 3 optimizes against the
                # currently active Phase II limits. Otherwise Phase I is used.
                # -----------------------------------------------------
                if enable_phase2:
                    active_lab_lcl = float(lab_after_lcl)
                    active_lab_ucl = float(lab_after_ucl)
                    active_line_lsl = float(line_after_lcl)
                    active_line_usl = float(line_after_ucl)
                    active_limit_name = "Phase II"
                else:
                    active_lab_lcl = float(lab_before_lcl)
                    active_lab_ucl = float(lab_before_ucl)
                    active_line_lsl = float(line_before_lcl)
                    active_line_usl = float(line_before_ucl)
                    active_limit_name = "Phase I"

                current_lab_center = (active_lab_lcl + active_lab_ucl) / 2.0
                line_target_center = (active_line_lsl + active_line_usl) / 2.0
                current_lab_half_width = (active_lab_ucl - active_lab_lcl) / 2.0

                # -----------------------------------------------------
                # ONE COIL = ONE OBSERVATION
                # -----------------------------------------------------
                task3_cols = [coil_col, lab_gloss_col, 'LINE_Gloss', 'Is_Phase_II']
                df_task3_source = df_filtered[task3_cols].copy()
                df_task3_source = df_task3_source.dropna(
                    subset=[coil_col, lab_gloss_col, 'LINE_Gloss']
                )

                df_task3 = (
                    df_task3_source
                    .groupby(coil_col, as_index=False)
                    .agg({
                        lab_gloss_col: 'mean',
                        'LINE_Gloss': 'mean',
                        'Is_Phase_II': 'first'
                    })
                )

                # Evaluate all historical observations against the ACTIVE limits.
                # This keeps Task 3 on one common target even if Phase I and II differ.
                df_task3['ACTIVE_LINE_PASS'] = (
                    (df_task3['LINE_Gloss'] >= active_line_lsl) &
                    (df_task3['LINE_Gloss'] <= active_line_usl)
                )
                df_task3['ACTIVE_LAB_ACCEPT'] = (
                    (df_task3[lab_gloss_col] >= active_lab_lcl) &
                    (df_task3[lab_gloss_col] <= active_lab_ucl)
                )

                def wilson_lower_bound(successes, n, confidence=0.95):
                    if n <= 0:
                        return np.nan
                    z = stats.norm.ppf(1 - (1 - confidence) / 2)
                    p_hat = successes / n
                    denominator = 1 + (z ** 2 / n)
                    centre = p_hat + (z ** 2 / (2 * n))
                    margin = z * np.sqrt(
                        (p_hat * (1 - p_hat) / n) +
                        (z ** 2 / (4 * n ** 2))
                    )
                    return (centre - margin) / denominator

                def fmt_range(low, high):
                    return f"{low:.1f} ~ {high:.1f}"

                # -----------------------------------------------------
                # BASELINE / CENTER METRICS
                # -----------------------------------------------------
                n_total_task3 = len(df_task3)
                historical_lab_min = (
                    float(df_task3[lab_gloss_col].min()) if n_total_task3 > 0 else np.nan
                )
                historical_lab_max = (
                    float(df_task3[lab_gloss_col].max()) if n_total_task3 > 0 else np.nan
                )

                current_control_df = df_task3[df_task3['ACTIVE_LAB_ACCEPT']].copy()
                current_n = len(current_control_df)

                # If no historical points fall inside the active LAB limits,
                # keep center optimization visible but do not invent performance evidence.
                if current_n > 0:
                    current_successes = int(current_control_df['ACTIVE_LINE_PASS'].sum())
                    current_pass_rate = current_successes / current_n
                    current_lcb = wilson_lower_bound(current_successes, current_n)
                    current_line_sd = (
                        current_control_df['LINE_Gloss'].std(ddof=1)
                        if current_n > 1 else 0.0
                    )
                    current_line_center = float(current_control_df['LINE_Gloss'].mean())
                    current_line_median = float(current_control_df['LINE_Gloss'].median())
                    line_center_error = line_target_center - current_line_center
                    current_center_shift = current_line_center - line_target_center
                    current_center_dev = abs(line_center_error)
                    current_safety_margin = min(
                        current_line_center - active_line_lsl,
                        active_line_usl - current_line_center
                    )
                    current_actual_lab_mean = float(current_control_df[lab_gloss_col].mean())
                else:
                    current_successes = 0
                    current_pass_rate = np.nan
                    current_lcb = np.nan
                    current_line_sd = np.nan
                    current_line_center = np.nan
                    current_line_median = np.nan
                    line_center_error = np.nan
                    current_center_shift = np.nan
                    current_center_dev = np.nan
                    current_safety_margin = np.nan
                    current_actual_lab_mean = np.nan

                unique_lab_values = df_task3[lab_gloss_col].nunique()

                # Historical LAB → LINE model.
                if (
                    n_total_task3 >= 3 and
                    unique_lab_values >= 2 and
                    df_task3[lab_gloss_col].std(ddof=1) > 0 and
                    df_task3['LINE_Gloss'].std(ddof=1) > 0
                ):
                    pearson_r, pearson_p = stats.pearsonr(
                        df_task3[lab_gloss_col],
                        df_task3['LINE_Gloss']
                    )
                    task3_r2 = pearson_r ** 2
                    slope, intercept, reg_r, reg_p, reg_stderr = stats.linregress(
                        df_task3[lab_gloss_col],
                        df_task3['LINE_Gloss']
                    )
                else:
                    pearson_r = np.nan
                    pearson_p = np.nan
                    task3_r2 = np.nan
                    slope = np.nan
                    intercept = np.nan
                    reg_stderr = np.nan

                overall_pass_rate = (
                    df_task3['ACTIVE_LINE_PASS'].mean() * 100
                    if n_total_task3 > 0 else np.nan
                )

                regression_reliable = (
                    pd.notna(task3_r2) and
                    pd.notna(slope) and
                    pd.notna(intercept) and
                    task3_r2 >= task3_reverse_min_r2 and
                    abs(slope) >= task3_reverse_min_abs_slope
                )

                # -----------------------------------------------------
                # CENTER-TO-CENTER OPTIMIZATION CORE
                # -----------------------------------------------------
                required_lab_shift = np.nan
                optimized_lab_center = np.nan
                optimized_lab_lcl = np.nan
                optimized_lab_ucl = np.nan
                regression_reverse_target = np.nan

                if pd.notna(line_center_error) and regression_reliable:
                    # Primary logic requested:
                    # LINE Center Error = LINE Target Center - Current LINE Center
                    # Required LAB Shift = LINE Center Error / LAB→LINE slope
                    # Optimized LAB Center = Current LAB Control Center + Required LAB Shift
                    required_lab_shift = line_center_error / slope
                    optimized_lab_center = current_lab_center + required_lab_shift
                    optimized_lab_lcl = optimized_lab_center - current_lab_half_width
                    optimized_lab_ucl = optimized_lab_center + current_lab_half_width

                    # Secondary model check. This is shown for reference only.
                    regression_reverse_target = (line_target_center - intercept) / slope

                # -----------------------------------------------------
                # TOP KPI ROW
                # -----------------------------------------------------
                k1, k2, k3, k4, k5, k6 = st.columns(6)
                k1.metric("Coils", f"{n_total_task3}")
                k2.metric(
                    "LAB→LINE R²",
                    f"{task3_r2:.3f}" if pd.notna(task3_r2) else "N/A"
                )
                k3.metric("LINE Target Center", f"{line_target_center:.2f}")
                k4.metric(
                    "Current LINE Center",
                    f"{current_line_center:.2f}" if pd.notna(current_line_center) else "N/A"
                )
                k5.metric(
                    "LINE Center Error",
                    f"{line_center_error:+.2f}" if pd.notna(line_center_error) else "N/A",
                    help="Target LINE Center - Current LINE Center."
                )
                k6.metric(
                    "Required LAB Shift",
                    f"{required_lab_shift:+.2f}" if pd.notna(required_lab_shift) else "N/A",
                    help="LINE Center Error / LAB→LINE slope."
                )

                if pd.notna(historical_lab_min) and pd.notna(historical_lab_max):
                    st.info(
                        f"🔎 **Active limits: {active_limit_name}** | "
                        f"LINE: **{active_line_lsl:.1f} ~ {active_line_usl:.1f} GU** → Target Center: **{line_target_center:.2f} GU** | "
                        f"LAB: **{active_lab_lcl:.1f} ~ {active_lab_ucl:.1f} GU** → Control Center: **{current_lab_center:.2f} GU**  \n"
                        f"Historical LAB range: **{historical_lab_min:.1f} ~ {historical_lab_max:.1f} GU** | "
                        f"Overall LINE pass vs active limits: **{overall_pass_rate:.1f}%**"
                    )

                # -----------------------------------------------------
                # BUILD HISTORICAL CANDIDATE TABLE
                # -----------------------------------------------------
                candidate_df = pd.DataFrame()

                if n_total_task3 >= task3_min_coils and unique_lab_values >= 2:
                    grid_min = np.floor(historical_lab_min / task3_step) * task3_step
                    grid_max = np.ceil(historical_lab_max / task3_step) * task3_step
                    grid = np.arange(
                        grid_min,
                        grid_max + task3_step * 0.5,
                        task3_step
                    )
                    target_rate = task3_target_pass / 100.0
                    min_coverage_rate = task3_min_coverage / 100.0
                    candidate_rows = []

                    for i, lower in enumerate(grid[:-1]):
                        for upper in grid[i + 1:]:
                            sub = df_task3[
                                (df_task3[lab_gloss_col] >= lower) &
                                (df_task3[lab_gloss_col] <= upper)
                            ].copy()

                            n = len(sub)
                            if n < task3_min_coils:
                                continue

                            successes = int(sub['ACTIVE_LINE_PASS'].sum())
                            pass_rate = successes / n
                            wilson_lcb = wilson_lower_bound(successes, n)
                            line_std = (
                                sub['LINE_Gloss'].std(ddof=1)
                                if n > 1 else 0.0
                            )
                            line_center = float(sub['LINE_Gloss'].mean())
                            line_median = float(sub['LINE_Gloss'].median())
                            center_shift = line_center - line_target_center
                            center_dev = abs(center_shift)
                            safety_margin = min(
                                line_center - active_line_lsl,
                                active_line_usl - line_center
                            )
                            line_p05 = float(sub['LINE_Gloss'].quantile(0.05))
                            line_p95 = float(sub['LINE_Gloss'].quantile(0.95))
                            p90_width = line_p95 - line_p05
                            width = upper - lower
                            lab_center = (lower + upper) / 2.0
                            coverage = n / n_total_task3 if n_total_task3 > 0 else 0.0

                            meets_pass = pass_rate >= target_rate
                            meets_coverage = coverage >= min_coverage_rate
                            confidence_ok = (
                                wilson_lcb >= current_lcb
                                if task3_require_confidence and pd.notna(current_lcb)
                                else True
                            )
                            safety_eligible = (
                                meets_pass and meets_coverage and confidence_ok
                            )
                            centered_eligible = (
                                safety_eligible and
                                center_dev <= task3_max_center_dev
                            )

                            if center_dev <= 0.5:
                                center_status = "Centered"
                            elif center_shift > 0:
                                center_status = "Shifted High"
                            else:
                                center_status = "Shifted Low"

                            candidate_rows.append({
                                "LAB Lower": float(lower),
                                "LAB Upper": float(upper),
                                "LAB Center": float(lab_center),
                                "Width": float(width),
                                "Coils": int(n),
                                "Coverage": float(coverage),
                                "LINE Pass Rate": float(pass_rate),
                                "95% Pass LCB": float(wilson_lcb),
                                "LINE Center": float(line_center),
                                "LINE Median": float(line_median),
                                "LINE Target": float(line_target_center),
                                "Center Deviation": float(center_dev),
                                "Center Shift": float(center_shift),
                                "Center Status": center_status,
                                "Safety Margin": float(safety_margin),
                                "LINE SD": float(line_std),
                                "LINE P05": float(line_p05),
                                "LINE P95": float(line_p95),
                                "P90 Width": float(p90_width),
                                "Safety Eligible": bool(safety_eligible),
                                "Centered Eligible": bool(centered_eligible)
                            })

                    candidate_df = pd.DataFrame(candidate_rows)

                # =====================================================
                # LOGIC 1 — HISTORICAL SAFETY CONTROL
                # =====================================================
                st.markdown("### Logic 1 — Historical Safety Control")
                provisional_safe = None

                if candidate_df.empty:
                    logic1_status = "INSUFFICIENT"
                    logic1_reason = (
                        "Not enough historical data to establish a provisional LAB control range."
                    )
                else:
                    safe_candidates = candidate_df[
                        candidate_df["Safety Eligible"]
                    ].copy()

                    if safe_candidates.empty:
                        logic1_status = "NO_SAFE_RANGE"
                        logic1_reason = (
                            "No historical LAB range satisfies the selected LINE pass, coverage and confidence requirements."
                        )
                    else:
                        safe_candidates = safe_candidates.sort_values(
                            by=[
                                "Width",
                                "LINE Pass Rate",
                                "LINE SD",
                                "Safety Margin",
                                "Coverage"
                            ],
                            ascending=[False, False, True, False, False]
                        ).reset_index(drop=True)
                        provisional_safe = safe_candidates.iloc[0]
                        logic1_status = "SAFE_RANGE_FOUND"
                        logic1_reason = (
                            "Widest historical LAB band satisfying the selected safety and coverage requirements."
                        )

                if logic1_status == "SAFE_RANGE_FOUND":
                    s = provisional_safe
                    st.success(
                        f"✅ **Provisional Safe LAB Control: {s['LAB Lower']:.1f} ~ {s['LAB Upper']:.1f} GU**  \n"
                        f"LAB Center: **{s['LAB Center']:.2f} GU** | "
                        f"LINE Pass: **{s['LINE Pass Rate']*100:.1f}%** | "
                        f"Coverage: **{s['Coverage']*100:.1f}%** | "
                        f"LINE Center: **{s['LINE Center']:.2f} GU** | "
                        f"Center Deviation: **{s['Center Deviation']:.2f} GU** | "
                        f"LINE SD: **{s['LINE SD']:.2f} GU**"
                    )
                    st.caption(logic1_reason)
                elif logic1_status == "NO_SAFE_RANGE":
                    st.warning("🟡 **No Historical Safe LAB Range Found**")
                    st.write(logic1_reason)
                else:
                    st.warning("🟡 **Insufficient Evidence for Historical Safety Control**")
                    st.write(logic1_reason)

                # =====================================================
                # LOGIC 2 — CENTER-TO-CENTER OPTIMIZATION
                # =====================================================
                st.markdown("### Logic 2 — Center-to-Center Optimization")

                best_historical_center = None
                if not candidate_df.empty:
                    safe_candidates_all = candidate_df[
                        candidate_df["Safety Eligible"]
                    ].copy()
                    if not safe_candidates_all.empty:
                        best_historical_center = safe_candidates_all.sort_values(
                            by=[
                                "Center Deviation",
                                "LINE SD",
                                "Safety Margin",
                                "LINE Pass Rate",
                                "Coverage",
                                "Width"
                            ],
                            ascending=[True, True, False, False, False, False]
                        ).iloc[0]

                c1, c2, c3, c4 = st.columns(4)
                c1.metric("Current LAB Control Center", f"{current_lab_center:.2f} GU")
                c2.metric("LINE Target Center", f"{line_target_center:.2f} GU")
                c3.metric(
                    "Current LINE Center",
                    f"{current_line_center:.2f} GU" if pd.notna(current_line_center) else "N/A"
                )
                c4.metric(
                    "LINE Gap to Target",
                    f"{line_center_error:+.2f} GU" if pd.notna(line_center_error) else "N/A"
                )

                if current_n == 0:
                    logic2_status = "NO_CURRENT_EVIDENCE"
                    logic2_reason = (
                        "No historical coil falls inside the active LAB control range, so the current LINE center cannot be estimated."
                    )
                    st.warning("🟡 **Current center cannot be estimated from the active LAB control band.**")
                    st.write(logic2_reason)
                elif current_center_dev <= task3_max_center_dev:
                    logic2_status = "ALREADY_CENTERED"
                    logic2_reason = (
                        f"Current LINE center is already within ±{task3_max_center_dev:.2f} GU of the LINE target center."
                    )
                    st.success(
                        f"✅ **LINE is already sufficiently centered.**  \n"
                        f"Target Center: **{line_target_center:.2f} GU** | "
                        f"Current LINE Center: **{current_line_center:.2f} GU** | "
                        f"Deviation: **{current_center_dev:.2f} GU**"
                    )
                    st.caption(logic2_reason)
                elif regression_reliable:
                    logic2_status = "SHIFT_CALCULATED"
                    logic2_reason = (
                        "The LAB→LINE relationship passes the selected R² and slope thresholds, so the LINE center gap "
                        "can be translated into a quantitative LAB center adjustment."
                    )

                    st.success(
                        f"🎯 **Optimized LAB Center: {optimized_lab_center:.2f} GU**  \n"
                        f"Current LAB Center: **{current_lab_center:.2f} GU** | "
                        f"Required LAB Shift: **{required_lab_shift:+.2f} GU**  \n"
                        f"Model-Centered LAB Band (same width as current control): "
                        f"**{optimized_lab_lcl:.2f} ~ {optimized_lab_ucl:.2f} GU**"
                    )

                    st.markdown(
                        f"""
                        **Center calculation**

                        - `LINE Target Center = (LINE LSL + LINE USL) / 2`
                          = ({active_line_lsl:.2f} + {active_line_usl:.2f}) / 2
                          = **{line_target_center:.2f} GU**
                        - `Current LAB Control Center = (LAB LCL + LAB UCL) / 2`
                          = ({active_lab_lcl:.2f} + {active_lab_ucl:.2f}) / 2
                          = **{current_lab_center:.2f} GU**
                        - `Current LINE Center = Mean(LINE output inside current LAB control)`
                          = **{current_line_center:.2f} GU**
                        - `LINE Center Error = LINE Target Center − Current LINE Center`
                          = {line_target_center:.2f} − {current_line_center:.2f}
                          = **{line_center_error:+.2f} GU**
                        - `Required LAB Shift = LINE Center Error / LAB→LINE slope`
                          = {line_center_error:+.2f} / {slope:.4f}
                          = **{required_lab_shift:+.2f} GU**
                        - `Optimized LAB Center = Current LAB Control Center + Required LAB Shift`
                          = {current_lab_center:.2f} + ({required_lab_shift:+.2f})
                          = **{optimized_lab_center:.2f} GU**
                        """
                    )

                    st.caption(
                        f"Regression check: slope = {slope:.4f} LINE GU / LAB GU, "
                        f"R² = {task3_r2:.3f}. Direct reverse-regression target = "
                        f"{regression_reverse_target:.2f} GU (reference only)."
                    )
                else:
                    logic2_status = "MODEL_WEAK"
                    r2_text = f"{task3_r2:.3f}" if pd.notna(task3_r2) else "N/A"
                    slope_text = f"{slope:.4f}" if pd.notna(slope) else "N/A"
                    logic2_reason = (
                        f"A full LAB shift is not calculated because the historical LAB→LINE model is not reliable enough "
                        f"under the selected thresholds (R²={r2_text}, slope={slope_text})."
                    )
                    st.warning("🟡 **Quantitative LAB center shift is not reliable enough yet.**")
                    st.write(logic2_reason)

                if best_historical_center is not None:
                    st.info(
                        f"📌 **Best historical safe centering evidence:** LAB "
                        f"{best_historical_center['LAB Lower']:.1f} ~ {best_historical_center['LAB Upper']:.1f} GU "
                        f"(center {best_historical_center['LAB Center']:.2f}) produced LINE center "
                        f"{best_historical_center['LINE Center']:.2f} GU, "
                        f"{best_historical_center['Center Deviation']:.2f} GU from target."
                    )

                # =====================================================
                # LOGIC 3 — CONTROLLED PILOT TOWARD OPTIMIZED CENTER
                # =====================================================
                st.markdown("### Logic 3 — Controlled Pilot Toward Optimized LAB Center")

                adaptive_status = "NOT_NEEDED"
                adaptive_reason = ""
                next_pilot_target = np.nan
                pilot_low = np.nan
                pilot_high = np.nan
                planned_shift = 0.0
                shift_direction = "Hold"
                model_based = False

                if current_n == 0:
                    adaptive_status = "UNAVAILABLE"
                    adaptive_reason = (
                        "A pilot direction cannot be calculated because the current LINE center is unavailable."
                    )

                elif current_center_dev <= task3_max_center_dev:
                    adaptive_status = "NOT_NEEDED"
                    adaptive_reason = (
                        "Current LINE center is already inside the selected centering tolerance."
                    )

                elif regression_reliable and pd.notna(required_lab_shift):
                    # Move only part of the full required shift in one pilot.
                    planned_shift = float(np.clip(
                        required_lab_shift,
                        -task3_max_step,
                        task3_max_step
                    ))
                    next_pilot_target = current_lab_center + planned_shift

                    # Keep the first pilot close to observed history.
                    allowed_low = historical_lab_min - task3_max_pilot_outside_history
                    allowed_high = historical_lab_max + task3_max_pilot_outside_history
                    next_pilot_target = float(np.clip(
                        next_pilot_target,
                        allowed_low,
                        allowed_high
                    ))
                    planned_shift = next_pilot_target - current_lab_center

                    shift_direction = (
                        "Increase LAB" if planned_shift > 0
                        else "Decrease LAB" if planned_shift < 0
                        else "Hold"
                    )
                    model_based = True
                    adaptive_status = "MODEL_BASED_PILOT"
                    adaptive_reason = (
                        f"Full center correction requires {required_lab_shift:+.2f} GU, but the next pilot is limited "
                        f"to ±{task3_max_step:.2f} GU and constrained near the historical LAB range."
                    )

                elif pd.notna(slope) and abs(slope) >= task3_reverse_min_abs_slope and pd.notna(line_center_error):
                    # Low-R²: use only the observed slope direction, never a full reverse target.
                    direction_sign = np.sign(line_center_error / slope)
                    planned_shift = float(direction_sign * task3_max_step)
                    next_pilot_target = current_lab_center + planned_shift

                    allowed_low = historical_lab_min - task3_max_pilot_outside_history
                    allowed_high = historical_lab_max + task3_max_pilot_outside_history
                    next_pilot_target = float(np.clip(
                        next_pilot_target,
                        allowed_low,
                        allowed_high
                    ))
                    planned_shift = next_pilot_target - current_lab_center

                    shift_direction = (
                        "Increase LAB" if planned_shift > 0
                        else "Decrease LAB" if planned_shift < 0
                        else "Hold"
                    )
                    adaptive_status = "DIRECTION_ONLY_PILOT"
                    adaptive_reason = (
                        "R² is below the model-based threshold, so the app does not claim a full optimized LAB target. "
                        "It only uses the fitted slope sign to choose a conservative one-step pilot direction."
                    )
                else:
                    adaptive_status = "UNAVAILABLE"
                    adaptive_reason = (
                        "The LAB→LINE slope is too weak or unavailable, so a data-supported pilot direction cannot be determined."
                    )

                if adaptive_status in ["MODEL_BASED_PILOT", "DIRECTION_ONLY_PILOT"]:
                    pilot_low = next_pilot_target - task3_pilot_half_width
                    pilot_high = next_pilot_target + task3_pilot_half_width

                if adaptive_status == "NOT_NEEDED":
                    st.success("✅ **No Additional LAB Shift Required**")
                    st.write(adaptive_reason)
                elif adaptive_status == "MODEL_BASED_PILOT":
                    st.info(
                        f"🧪 **Next Pilot LAB Target: {next_pilot_target:.2f} GU**  \n"
                        f"Pilot Range: **{pilot_low:.2f} ~ {pilot_high:.2f} GU** | "
                        f"Action: **{shift_direction} {abs(planned_shift):.2f} GU**  \n"
                        f"Full Optimized LAB Center: **{optimized_lab_center:.2f} GU** | "
                        f"Full Required Shift: **{required_lab_shift:+.2f} GU** | "
                        f"R²: **{task3_r2:.3f}** | Slope: **{slope:.4f}**"
                    )
                    st.caption(adaptive_reason)
                elif adaptive_status == "DIRECTION_ONLY_PILOT":
                    st.warning(
                        f"🧪 **Conservative Direction-Only Pilot: {next_pilot_target:.2f} GU**  \n"
                        f"Pilot Range: **{pilot_low:.2f} ~ {pilot_high:.2f} GU** | "
                        f"Action: **{shift_direction} {abs(planned_shift):.2f} GU**  \n"
                        f"LINE Target Center: **{line_target_center:.2f} GU** | "
                        f"Current LINE Center: **{current_line_center:.2f} GU**"
                    )
                    st.caption(adaptive_reason)
                else:
                    st.error("🔴 **Pilot shift cannot be calculated reliably from the available data.**")
                    st.write(adaptive_reason)

                # =====================================================
                # FINAL RECOMMENDATION
                # =====================================================
                st.markdown("### Final Task 3 Control Recommendation")

                if logic2_status == "ALREADY_CENTERED":
                    st.success(
                        f"✅ **Keep Current LAB Control: {active_lab_lcl:.1f} ~ {active_lab_ucl:.1f} GU**  \n"
                        f"Current LAB Center: **{current_lab_center:.2f} GU** | "
                        f"LINE Target Center: **{line_target_center:.2f} GU** | "
                        f"Current LINE Center: **{current_line_center:.2f} GU**"
                    )
                elif adaptive_status == "MODEL_BASED_PILOT":
                    safe_text = (
                        f"{provisional_safe['LAB Lower']:.1f} ~ {provisional_safe['LAB Upper']:.1f} GU"
                        if provisional_safe is not None else "Not established"
                    )
                    st.warning(
                        f"🎯 **Center Optimization Target**  \n"
                        f"Current LAB Control Center: **{current_lab_center:.2f} GU**  \n"
                        f"Optimized LAB Center: **{optimized_lab_center:.2f} GU** "
                        f"(required shift {required_lab_shift:+.2f} GU)  \n"
                        f"Model-Centered LAB Band: **{optimized_lab_lcl:.2f} ~ {optimized_lab_ucl:.2f} GU**  \n"
                        f"Historical Safe LAB Range: **{safe_text}**  \n"
                        f"Next Pilot: **{next_pilot_target:.2f} GU** ({pilot_low:.2f} ~ {pilot_high:.2f} GU)"
                    )
                elif provisional_safe is not None:
                    st.info(
                        f"ℹ️ **Provisional Historical Safe LAB Range: "
                        f"{provisional_safe['LAB Lower']:.1f} ~ {provisional_safe['LAB Upper']:.1f} GU**  \n"
                        f"A final center-adjusted LAB target is not released because the LAB→LINE response model "
                        f"does not yet meet the selected reliability criteria."
                    )
                else:
                    st.error(
                        "🔴 **No reliable optimized LAB control target can currently be established from the available history.**"
                    )

                # =====================================================
                # COMPARISON SUMMARY TABLE
                # =====================================================
                st.markdown("#### Comparison Summary")

                summary_rows = [{
                    "Method": "Current Active LAB Control",
                    "LAB Range / Target": f"{active_lab_lcl:.1f} ~ {active_lab_ucl:.1f}",
                    "LAB Center": current_lab_center,
                    "LAB Shift": 0.0,
                    "Coils": current_n,
                    "Coverage (%)": (
                        current_n / n_total_task3 * 100
                        if n_total_task3 > 0 else np.nan
                    ),
                    "LINE Pass (%)": (
                        current_pass_rate * 100
                        if pd.notna(current_pass_rate) else np.nan
                    ),
                    "LINE Center": current_line_center,
                    "LINE Target": line_target_center,
                    "Center Deviation": current_center_dev,
                    "LINE SD": current_line_sd,
                    "Status": "Baseline"
                }]

                if provisional_safe is not None:
                    summary_rows.append({
                        "Method": "Provisional Safe Historical Range",
                        "LAB Range / Target": (
                            f"{provisional_safe['LAB Lower']:.1f} ~ "
                            f"{provisional_safe['LAB Upper']:.1f}"
                        ),
                        "LAB Center": float(provisional_safe["LAB Center"]),
                        "LAB Shift": float(provisional_safe["LAB Center"] - current_lab_center),
                        "Coils": int(provisional_safe["Coils"]),
                        "Coverage (%)": float(provisional_safe["Coverage"]) * 100,
                        "LINE Pass (%)": float(provisional_safe["LINE Pass Rate"]) * 100,
                        "LINE Center": float(provisional_safe["LINE Center"]),
                        "LINE Target": line_target_center,
                        "Center Deviation": float(provisional_safe["Center Deviation"]),
                        "LINE SD": float(provisional_safe["LINE SD"]),
                        "Status": "Historical Safety"
                    })

                if best_historical_center is not None:
                    summary_rows.append({
                        "Method": "Best Historical Centered Evidence",
                        "LAB Range / Target": (
                            f"{best_historical_center['LAB Lower']:.1f} ~ "
                            f"{best_historical_center['LAB Upper']:.1f}"
                        ),
                        "LAB Center": float(best_historical_center["LAB Center"]),
                        "LAB Shift": float(best_historical_center["LAB Center"] - current_lab_center),
                        "Coils": int(best_historical_center["Coils"]),
                        "Coverage (%)": float(best_historical_center["Coverage"]) * 100,
                        "LINE Pass (%)": float(best_historical_center["LINE Pass Rate"]) * 100,
                        "LINE Center": float(best_historical_center["LINE Center"]),
                        "LINE Target": line_target_center,
                        "Center Deviation": float(best_historical_center["Center Deviation"]),
                        "LINE SD": float(best_historical_center["LINE SD"]),
                        "Status": "Historical Centering Evidence"
                    })

                if pd.notna(optimized_lab_center):
                    summary_rows.append({
                        "Method": "Model-Centered LAB Target",
                        "LAB Range / Target": f"{optimized_lab_lcl:.2f} ~ {optimized_lab_ucl:.2f}",
                        "LAB Center": optimized_lab_center,
                        "LAB Shift": required_lab_shift,
                        "Coils": np.nan,
                        "Coverage (%)": np.nan,
                        "LINE Pass (%)": np.nan,
                        "LINE Center": np.nan,
                        "LINE Target": line_target_center,
                        "Center Deviation": np.nan,
                        "LINE SD": np.nan,
                        "Status": "Calculated Center Target"
                    })

                if adaptive_status in ["MODEL_BASED_PILOT", "DIRECTION_ONLY_PILOT"]:
                    summary_rows.append({
                        "Method": "Next Pilot",
                        "LAB Range / Target": f"{next_pilot_target:.2f} ({pilot_low:.2f}~{pilot_high:.2f})",
                        "LAB Center": next_pilot_target,
                        "LAB Shift": planned_shift,
                        "Coils": np.nan,
                        "Coverage (%)": np.nan,
                        "LINE Pass (%)": np.nan,
                        "LINE Center": np.nan,
                        "LINE Target": line_target_center,
                        "Center Deviation": np.nan,
                        "LINE SD": np.nan,
                        "Status": (
                            "Model-Based Pilot" if model_based
                            else "Direction-Only Pilot"
                        )
                    })

                summary_df = pd.DataFrame(summary_rows)
                for c in [
                    "LAB Center",
                    "LAB Shift",
                    "Coverage (%)",
                    "LINE Pass (%)",
                    "LINE Center",
                    "LINE Target",
                    "Center Deviation",
                    "LINE SD"
                ]:
                    if c in summary_df.columns:
                        summary_df[c] = summary_df[c].round(2)

                st.dataframe(
                    summary_df,
                    use_container_width=True,
                    hide_index=True
                )

                # =====================================================
                # LAB → LINE RELATIONSHIP CHART
                # =====================================================
                st.markdown("#### LAB → LINE Relationship")

                fig_opt, ax_opt = plt.subplots(figsize=(12, 6))
                ax_opt.set_facecolor('#f2f2f2')

                pass_mask = df_task3['ACTIVE_LINE_PASS']
                fail_mask = ~df_task3['ACTIVE_LINE_PASS']

                ax_opt.scatter(
                    df_task3.loc[pass_mask, lab_gloss_col],
                    df_task3.loc[pass_mask, 'LINE_Gloss'],
                    marker='o', s=65, alpha=0.8, label='LINE Pass'
                )

                if fail_mask.any():
                    ax_opt.scatter(
                        df_task3.loc[fail_mask, lab_gloss_col],
                        df_task3.loc[fail_mask, 'LINE_Gloss'],
                        marker='X', s=85, label='LINE Outlier'
                    )

                # Active LAB limits and center.
                ax_opt.axvline(
                    active_lab_lcl,
                    linestyle='--', linewidth=1.5,
                    label=f'Active LAB Limits ({active_limit_name})'
                )
                ax_opt.axvline(active_lab_ucl, linestyle='--', linewidth=1.5)
                ax_opt.axvline(
                    current_lab_center,
                    linestyle=':', linewidth=1.8,
                    label='Current LAB Control Center'
                )

                # Active LINE limits and exact target midpoint.
                ax_opt.axhline(
                    active_line_lsl,
                    linestyle=':', linewidth=1.8,
                    label=f'Active LINE Limits ({active_limit_name})'
                )
                ax_opt.axhline(active_line_usl, linestyle=':', linewidth=1.8)
                ax_opt.axhline(
                    line_target_center,
                    linestyle='-.', linewidth=2.0,
                    label=f'LINE Target Center ({line_target_center:.2f})'
                )

                if pd.notna(current_line_center):
                    ax_opt.axhline(
                        current_line_center,
                        linestyle='--', linewidth=1.3,
                        label=f'Current LINE Center ({current_line_center:.2f})'
                    )

                if pd.notna(slope):
                    x_reg = np.linspace(
                        df_task3[lab_gloss_col].min(),
                        df_task3[lab_gloss_col].max(),
                        100
                    )
                    y_reg = intercept + slope * x_reg
                    ax_opt.plot(
                        x_reg,
                        y_reg,
                        linewidth=1.8,
                        label=f'Linear Fit (R²={task3_r2:.3f})'
                    )

                if provisional_safe is not None:
                    ax_opt.axvspan(
                        float(provisional_safe["LAB Lower"]),
                        float(provisional_safe["LAB Upper"]),
                        alpha=0.10,
                        label='Provisional Safe LAB Range'
                    )

                if pd.notna(optimized_lab_center):
                    ax_opt.axvline(
                        optimized_lab_center,
                        linestyle='-.', linewidth=2.2,
                        label=f'Optimized LAB Center ({optimized_lab_center:.2f})'
                    )
                    ax_opt.axvspan(
                        optimized_lab_lcl,
                        optimized_lab_ucl,
                        alpha=0.08,
                        label='Model-Centered LAB Band'
                    )

                if adaptive_status in ["MODEL_BASED_PILOT", "DIRECTION_ONLY_PILOT"]:
                    ax_opt.axvline(
                        next_pilot_target,
                        linestyle='--', linewidth=2.0,
                        label=f'Next Pilot LAB Target ({next_pilot_target:.2f})'
                    )

                ax_opt.set_title(
                    f"LAB Input vs LINE Output: {selected_paint}",
                    fontsize=15,
                    fontweight="bold",
                    pad=18
                )
                ax_opt.set_xlabel("LAB Input Gloss (GU)")
                ax_opt.set_ylabel("LINE Output Gloss (GU)")
                ax_opt.grid(True, alpha=0.35)
                ax_opt.legend(
                    bbox_to_anchor=(1.02, 1),
                    loc="upper left",
                    frameon=True,
                    edgecolor="black"
                )
                fig_opt.subplots_adjust(right=0.76)
                st.pyplot(fig_opt)

                # =====================================================
                # HISTORICAL CANDIDATE TABLE
                # =====================================================
                st.markdown("#### Historical Candidate LAB Range Evaluation")

                if not candidate_df.empty:
                    candidate_view = candidate_df.copy()
                    candidate_view["LAB Range"] = candidate_view.apply(
                        lambda r: fmt_range(r["LAB Lower"], r["LAB Upper"]),
                        axis=1
                    )
                    candidate_view["Coverage (%)"] = (
                        candidate_view["Coverage"] * 100
                    ).round(1)
                    candidate_view["LINE Pass (%)"] = (
                        candidate_view["LINE Pass Rate"] * 100
                    ).round(1)
                    candidate_view["95% Pass LCB (%)"] = (
                        candidate_view["95% Pass LCB"] * 100
                    ).round(1)

                    for c in [
                        "LAB Center",
                        "Width",
                        "LINE Center",
                        "LINE Median",
                        "LINE Target",
                        "Center Deviation",
                        "Center Shift",
                        "Safety Margin",
                        "LINE SD",
                        "LINE P05",
                        "LINE P95",
                        "P90 Width"
                    ]:
                        candidate_view[c] = candidate_view[c].round(2)

                    candidate_view["Safety Decision"] = np.where(
                        candidate_view["Safety Eligible"],
                        "Safe",
                        "Not Safe"
                    )
                    candidate_view["Center Decision"] = np.where(
                        candidate_view["Centered Eligible"],
                        "Centered",
                        "Not Centered"
                    )

                    candidate_view = candidate_view.sort_values(
                        by=[
                            "Safety Eligible",
                            "Centered Eligible",
                            "Center Deviation",
                            "LINE SD",
                            "Safety Margin",
                            "LINE Pass (%)",
                            "Coverage (%)",
                            "Width"
                        ],
                        ascending=[False, False, True, True, False, False, False, False]
                    )

                    show_cols = [
                        "LAB Range",
                        "LAB Center",
                        "Width",
                        "Coils",
                        "Coverage (%)",
                        "LINE Pass (%)",
                        "LINE Center",
                        "LINE Median",
                        "LINE Target",
                        "Center Deviation",
                        "Center Shift",
                        "Center Status",
                        "Safety Margin",
                        "LINE SD",
                        "LINE P05",
                        "LINE P95",
                        "P90 Width",
                        "95% Pass LCB (%)",
                        "Safety Decision",
                        "Center Decision"
                    ]

                    st.dataframe(
                        candidate_view[show_cols].head(30),
                        use_container_width=True,
                        hide_index=True
                    )
                else:
                    st.info("Historical candidate scan could not be completed.")

                # =====================================================
                # METHOD SUMMARY
                # =====================================================
                with st.expander("ℹ️ Task 3 Method"):
                    st.markdown(
                        f"""
                        **Three-logics control method**

                        **Logic 1 — Historical Safety Control**
                        - Scan historical LAB input bands.
                        - Keep candidate bands that satisfy LINE pass rate, minimum historical coverage and confidence requirements.
                        - The selected provisional range is a historical safety guard band; it is not automatically the centered target.

                        **Logic 2 — Center-to-Center Optimization**
                        - Active LINE target center is fixed at the midpoint of the active LINE limits:
                          **({active_line_lsl:.2f} + {active_line_usl:.2f}) / 2 = {line_target_center:.2f} GU**.
                        - Current LAB control center is fixed at the midpoint of the active LAB limits:
                          **({active_lab_lcl:.2f} + {active_lab_ucl:.2f}) / 2 = {current_lab_center:.2f} GU**.
                        - Current LINE center is the mean LINE output of historical coils whose LAB values fall inside the active LAB control range.
                        - LINE Center Error = LINE Target Center − Current LINE Center.
                        - When the LAB→LINE model is reliable:
                          **Required LAB Shift = LINE Center Error / slope**.
                        - Optimized LAB Center = Current LAB Control Center + Required LAB Shift.
                        - The model-centered LAB band keeps the same width as the current LAB control band and shifts its center toward the calculated optimum.

                        **Logic 3 — Controlled Pilot Toward Optimized LAB Center**
                        - Do not jump directly to the full optimized center when the required movement is large.
                        - Limit each pilot to **±{task3_max_step:.2f} GU** and keep it near the historical LAB range.
                        - If R² is below the full-model threshold but slope direction remains usable, only a conservative direction-only pilot is allowed.
                        - If slope itself is too weak, Task 3 does not invent a LAB adjustment.
                        - After each pilot, append the new LAB + LINE result and rerun Task 3 so the center estimate and slope are updated.
                        """
                    )


            # =========================================================
            # DATA TABLE SUMMARY (Always visible)
            # =========================================================
            with st.expander("📋 View Coil Data Details"):
                display_df = df_filtered[[date_col, 'Batch_Input_Date', batch_col, coil_col, 'Is_Phase_II', lab_gloss_col, line_north_col, line_south_col, 'LINE_Gloss']].copy()
                display_df['Batch_Input_Date'] = display_df['Batch_Input_Date'].dt.strftime('%Y-%m-%d')
                display_df.columns = ["Coil Prod Date", "Batch Inspect Date", "Batch Number", "Coil ID", "Uses New Limits?", "LAB Input", "LINE North", "LINE South", "LINE Avg"]
                display_df['LINE Avg'] = display_df['LINE Avg'].round(2)
                st.dataframe(display_df, use_container_width=True)
else:
    st.warning("No data found or file is missing. Please check the directory path.")
